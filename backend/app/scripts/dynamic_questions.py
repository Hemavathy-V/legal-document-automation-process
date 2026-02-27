"""
Template Parser Module
Extracts placeholders from DOCX contract templates
Provides reusable contract processing utilities for API workflows
"""

import re
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from typing import Dict, List, Any, Tuple, Optional

# Add backend to path for logger import
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..', '..'))
from backend.app.core.logger import get_logger

logger = get_logger(__name__)


class TemplateParser:
    """Parse DOCX templates and extract placeholders with proper handling of loops"""
    
    def __init__(self, template_path: Path):
        """
        Initialize parser with a template file
        
        Args:
            template_path: Path to the DOCX template file
        """
        self.template_path = Path(template_path)
        logger.debug(f"Initializing TemplateParser for: {self.template_path}")
        self.template_text = self._load_template()
        self.simple_placeholders = []
        self.loop_placeholders = {}
        self._extract_all_placeholders()
        logger.info(
            f"Template parsed: {self.template_path.name} | "
            f"simple_placeholders={len(self.simple_placeholders)} | "
            f"loop_placeholders={len(self.loop_placeholders)}"
        )
    
    def _load_template(self) -> str:
        """Load and extract text from DOCX template"""
        logger.debug(f"Loading DOCX template: {self.template_path}")
        doc = Document(self.template_path)
        # Extract text from all paragraphs
        full_text = "\n".join([para.text for para in doc.paragraphs])
        logger.debug(f"Loaded {len(doc.paragraphs)} paragraphs from template: {self.template_path.name}")
        return full_text
    
    def _extract_all_placeholders(self) -> None:
        """Extract both simple and loop placeholders from template"""
        logger.debug(f"Extracting placeholders from template text: {self.template_path.name}")
        # Extract loop placeholders: {{#fieldName}}...{{/fieldName}}
        loop_pattern = r"\{\{#(\w+)\}\}(.*?)\{\{/\1\}\}"
        loop_matches = re.findall(loop_pattern, self.template_text, re.DOTALL)
        
        for loop_field, loop_content in loop_matches:
            # Extract nested fields within the loop
            nested_pattern = r"\{\{(\w+)\}\}"
            nested_fields = re.findall(nested_pattern, loop_content)
            self.loop_placeholders[loop_field] = nested_fields
            logger.debug(f"Detected loop placeholder '{loop_field}' with {len(nested_fields)} nested fields")
        
        # Extract simple placeholders: {{fieldName}}
        # But exclude those already in loops
        simple_pattern = r"\{\{([\w.]+)\}\}"
        simple_matches = re.findall(simple_pattern, self.template_text)
        
        # Filter out loop declarations
        for match in simple_matches:
            if match not in self.loop_placeholders:
                if match not in self.simple_placeholders:
                    self.simple_placeholders.append(match)
        logger.debug(
            f"Placeholder extraction complete for {self.template_path.name}: "
            f"simple={len(self.simple_placeholders)}, loops={len(self.loop_placeholders)}"
        )
    
    def get_simple_placeholders(self) -> List[str]:
        """Get list of simple placeholders"""
        return self.simple_placeholders
    
    def get_loop_placeholders(self) -> Dict[str, List[str]]:
        """Get dictionary of loop placeholders with their nested fields"""
        return self.loop_placeholders
    
    def get_all_placeholders(self) -> Tuple[List[str], Dict[str, List[str]]]:
        """Get both simple and loop placeholders"""
        return self.simple_placeholders, self.loop_placeholders
    
    def get_template_filename(self) -> str:
        """Get template filename without extension"""
        return self.template_path.stem.replace("-Template", "").replace("-", "")


class ContractProcessor:
    """Processor for template lookup, placeholder extraction, and JSON output generation"""
    
    def __init__(self, templates_dir: Path, output_dir: Path):
        """
        Initialize processor
        
        Args:
            templates_dir: Directory containing DOCX templates
            output_dir: Directory for saving JSON outputs
        """
        self.templates_dir = Path(templates_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"ContractProcessor initialized | templates_dir={self.templates_dir} | output_dir={self.output_dir}")
    
    def list_available_templates(self) -> List[Path]:
        """List all available DOCX templates"""
        templates = list(self.templates_dir.glob("*.docx"))
        logger.debug(f"Found {len(templates)} template(s) in {self.templates_dir}")
        return sorted(templates)

    def list_template_names(self) -> List[str]:
        """List available template names without extensions"""
        names = [template.stem for template in self.list_available_templates()]
        logger.debug(f"Template names resolved: {names}")
        return names

    def get_template_path(self, template_name: str) -> Optional[Path]:
        """Resolve a template name to a DOCX path"""
        normalized_name = template_name.strip()
        if not normalized_name:
            logger.warning("Empty template name received while resolving template path")
            return None

        # Convert spaces to hyphens to match file naming convention
        normalized_name = normalized_name.replace(" ", "-")
        template_file = normalized_name if normalized_name.lower().endswith(".docx") else f"{normalized_name}.docx"
        template_path = self.templates_dir / template_file
        if template_path.exists():
            logger.debug(f"Template path resolved: {template_path}")
            return template_path

        logger.warning(f"Template not found: {template_file} in {self.templates_dir}")
        return None

    def get_template_placeholders(self, template_name: str) -> Tuple[List[str], Dict[str, List[str]]]:
        """Get simple and loop placeholders for a specific template"""
        logger.info(f"Fetching placeholders for template: {template_name}")
        template_path = self.get_template_path(template_name)
        if not template_path:
            logger.error(f"Cannot fetch placeholders. Template not found: {template_name}")
            raise FileNotFoundError(f"Template '{template_name}' not found")

        parser = TemplateParser(template_path)
        logger.info(f"Placeholders extracted successfully for template: {template_name}")
        return parser.get_all_placeholders()

    def save_contract_data(self, template_name: str, data: Dict[str, Any]) -> Path:
        """Save contract payload to output directory as JSON"""
        logger.info(f"Saving contract data for template: {template_name}")
        template_path = self.get_template_path(template_name)
        if not template_path:
            logger.error(f"Cannot save contract data. Template not found: {template_name}")
            raise FileNotFoundError(f"Template '{template_name}' not found")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{template_path.stem}_{timestamp}.json"
        filepath = self.output_dir / filename

        with open(filepath, "w", encoding="utf-8") as file_obj:
            json.dump(data, file_obj, indent=4, ensure_ascii=False)

        logger.info(
            f"Contract data saved: {filepath} | "
            f"fields_count={self._count_fields(data)}"
        )
        return filepath

    @staticmethod
    def _count_fields(data: Dict) -> int:
        """Count total number of fields in collected data"""
        count = 0
        for value in data.values():
            if isinstance(value, list):
                count += sum(len(item) if isinstance(item, dict) else 1 for item in value)
            else:
                count += 1
        logger.debug(f"Computed fields count: {count}")
        return count
